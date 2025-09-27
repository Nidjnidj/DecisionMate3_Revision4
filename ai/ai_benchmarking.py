# ai_benchmarking.py
"""
DecisionMate — AI Benchmarking
Compare your project's cost/schedule/economics/procurement to similar projects & components.

What it does
- Loads your benchmark datasets (projects/components/procurement, optional FX table).
- Uses AI to propose peer-set filters from your project context (complexity & scope).
- Finds top-N similar projects (sector/region/complexity/scope/tags).
- Benchmarks components (pump, compressor, pipe, etc.) vs. peer distributions; flags high/low.
- NEW: Procurement lead-time & price benchmarking vs peers by package.
- NEW: Vendor mix heatmap (peer vendor distribution across packages; optional overlay with your vendors).
- Produces AI review summary & recommendations; exports CSV/DOCX; saves as artifact.

CSV Schemas (minimal)
  projects.csv:
    project_id, company, sector, region, complexity, scope_tags,
    total_capex_usd, total_duration_days, start_year, currency?, source_url?
  components.csv:
    project_id, component_type, spec_size, count, unit_cost, install_cost,
    duration_days, vendor?, currency?
  procurement.csv (optional peers):
    project_id, package, vendor, lead_time_days?, price?, currency?
  my_components.csv (your current project items to compare):
    component_type, spec_size, count, unit_cost, install_cost, duration_days, vendor?, currency?
  my_procurement.csv (optional, your current procurement):
    package, vendor, lead_time_days?, price?, currency?

  fx.csv (optional currency normalization):
    currency, to_usd_rate  # e.g., EUR,1.10 ; GBP,1.25

Install (if needed):
    pip install plotly python-docx
"""

from __future__ import annotations
import io, json, math, uuid
from dataclasses import dataclass
from typing import Optional, Dict, List, Tuple

import streamlit as st
import pandas as pd
import numpy as np

# Optional DOCX
try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# Plotly charts (optional)
try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_OK = True
except Exception:
    PLOTLY_OK = False

# App services
from .providers import get_chat_callable
from .tools import collect_phase_context, format_context_text

# -----------------------------
# Namespaced keys
# -----------------------------
if "AI_BENCH_NS" not in st.session_state:
    st.session_state["AI_BENCH_NS"] = f"bench_{uuid.uuid4().hex[:8]}"

def k(group: str, name: str) -> str:
    return f"{st.session_state['AI_BENCH_NS']}::{group}::{name}"

# -----------------------------
# Helper utilities
# -----------------------------
def _docx_bytes(title: str, body: str) -> Optional[bytes]:
    if not DOCX_OK:
        return None
    try:
        doc = Document()
        doc.add_heading(title or "AI Benchmarking Review", level=1)
        for line in (body or "").splitlines():
            if line.strip().startswith("#"):
                doc.add_heading(line.strip("# ").strip(), level=2)
            else:
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None

def _save_artifact_button(artifact_type: str, payload: dict, pid: str, ph: str, key: str):
    if st.button("Save as Artifact", key=key):
        try:
            from artifact_registry import save
            save(
                project_id=pid,
                phase_id=ph,
                artifact_type=artifact_type,
                data=payload,
                status="Draft",
            )
            st.success(f"Saved artifact: {artifact_type} (Draft)")
        except Exception as e:
            st.error(f"Could not save artifact: {e}")

def normalize_currency(df: pd.DataFrame, fx: pd.DataFrame, value_cols: List[str], currency_col: str = "currency") -> pd.DataFrame:
    """Convert value columns to USD using fx table currency->to_usd_rate. Leaves values if currency missing."""
    if df is None or df.empty or fx is None or fx.empty or currency_col not in df.columns:
        return df
    fx_map = dict(zip(fx["currency"], fx["to_usd_rate"]))
    out = df.copy()
    out[currency_col] = out[currency_col].fillna("USD")
    rates = out[currency_col].map(lambda c: float(fx_map.get(str(c), 1.0)))
    for c in value_cols:
        if c in out.columns:
            out[c] = pd.to_numeric(out[c], errors="coerce")
            out[c] = out[c] * rates
    out["currency"] = "USD"
    return out

def parse_tags(s: str) -> set:
    if not isinstance(s, str):
        return set()
    return set([x.strip().lower() for x in s.split(",") if x.strip()])

def complexity_score(v: str) -> int:
    if not isinstance(v, str):
        return 1
    t = v.strip().lower()
    return {"low":0, "medium":1, "med":1, "high":2, "very high":3, "very_high":3}.get(t, 1)

def jaccard(a: set, b: set) -> float:
    if not a and not b:
        return 1.0
    if not a or not b:
        return 0.0
    return len(a & b) / float(len(a | b))

def similarity_row(row, target_sector, target_region, target_complexity, target_tags_set):
    s_sector = 1.0 if str(row.get("sector","")).strip().lower() == target_sector else 0.0
    s_region = 1.0 if str(row.get("region","")).strip().lower() == target_region else 0.0
    s_comp = 1.0 - min(1.0, abs(complexity_score(row.get("complexity")) - complexity_score(target_complexity)) / 3.0)
    s_tags  = jaccard(parse_tags(row.get("scope_tags","")), target_tags_set)
    return 0.35*s_sector + 0.25*s_region + 0.25*s_comp + 0.15*s_tags

def ai_propose_filters(context_text: str) -> dict:
    """Ask LLM to suggest sector/region/complexity and scope tags for peer-set selection."""
    llm = get_chat_callable()
    system = (
        "You are a capital projects benchmarking analyst. "
        "From the given project context, infer sector, region, complexity (low/medium/high), "
        "and a short comma-separated list of scope tags. Return strict JSON."
    )
    user = f"""
PROJECT CONTEXT (summarized):
{context_text[:5000]}

Return JSON like:
{{
  "sector": "oil & gas upstream",
  "region": "middle east",
  "complexity": "high",
  "scope_tags": "pump, compressor, pipeline, greenfield"
}}
"""
    try:
        raw = llm(user, system)
        s = str(raw).strip()
        start, end = s.find("{"), s.rfind("}")
        js = s[start:end+1] if start != -1 and end != -1 else "{}"
        data = json.loads(js) if js else {}
        return {
            "sector": (data.get("sector") or "").strip().lower(),
            "region": (data.get("region") or "").strip().lower(),
            "complexity": (data.get("complexity") or "").strip().lower() or "medium",
            "scope_tags": (data.get("scope_tags") or "").strip().lower(),
        }
    except Exception:
        return {"sector":"", "region":"", "complexity":"medium", "scope_tags":""}

def compare_value_to_peer(value: float, peer_series: pd.Series) -> dict:
    """Return pct difference vs P50 and whether high/low/outlier."""
    s = pd.to_numeric(peer_series, errors="coerce").dropna()
    if s.empty or value is None or not np.isfinite(value):
        return {"status":"n/a", "pct_vs_p50": np.nan, "p50": np.nan, "p25": np.nan, "p75": np.nan}
    p25, p50, p75, p90 = np.percentile(s, [25,50,75,90])
    pct = (float(value) - p50) / (p50 if p50 != 0 else 1.0)
    status = "about average"
    if value > p75: status = "higher than peers"
    elif value < p25: status = "lower than peers"
    return {"status": status, "pct_vs_p50": pct, "p50": p50, "p25": p25, "p75": p75, "p90": p90}

# -----------------------------
# Procurement analytics (NEW)
# -----------------------------
def procurement_peer_stats(proc_df: pd.DataFrame, peer_ids: set) -> Tuple[pd.DataFrame, pd.DataFrame]:
    """
    Return (lead_time_stats, price_stats) for peers by package.
    Each has columns: package, n, p25, p50, p75, p90, mean, min, max
    """
    df = proc_df.copy()
    df = df[df["project_id"].astype(str).isin({str(x) for x in peer_ids})]
    out = {}
    for metric in ["lead_time_days", "price"]:
        if metric in df.columns:
            g = df.groupby("package")[metric].apply(lambda s: pd.Series({
                "n": s.dropna().shape[0],
                "p25": np.nanpercentile(s.dropna(), 25) if s.dropna().shape[0] else np.nan,
                "p50": np.nanpercentile(s.dropna(), 50) if s.dropna().shape[0] else np.nan,
                "p75": np.nanpercentile(s.dropna(), 75) if s.dropna().shape[0] else np.nan,
                "p90": np.nanpercentile(s.dropna(), 90) if s.dropna().shape[0] else np.nan,
                "mean": s.dropna().mean(),
                "min": s.dropna().min() if s.dropna().shape[0] else np.nan,
                "max": s.dropna().max() if s.dropna().shape[0] else np.nan,
            })).reset_index()
        else:
            g = pd.DataFrame(columns=["package","n","p25","p50","p75","p90","mean","min","max"])
        out[metric] = g
    return out.get("lead_time_days"), out.get("price")

def compare_my_procurement(my_proc: pd.DataFrame, peer_lead: pd.DataFrame, peer_price: pd.DataFrame) -> pd.DataFrame:
    """
    Join your packages to peer stats and compute status vs p50 for lead_time_days and price.
    """
    if my_proc is None or my_proc.empty:
        return pd.DataFrame()
    mine = my_proc.copy()
    for c in ["lead_time_days","price"]:
        if c in mine.columns:
            mine[c] = pd.to_numeric(mine[c], errors="coerce")
    lead = peer_lead.rename(columns={"p50":"lead_p50","p25":"lead_p25","p75":"lead_p75","p90":"lead_p90"}) if peer_lead is not None else pd.DataFrame()
    price = peer_price.rename(columns={"p50":"price_p50","p25":"price_p25","p75":"price_p75","p90":"price_p90"}) if peer_price is not None else pd.DataFrame()
    out = mine.merge(lead[["package","lead_p25","lead_p50","lead_p75","lead_p90","n"]], on="package", how="left")
    if not price.empty:
        out = out.merge(price[["package","price_p25","price_p50","price_p75","price_p90","n"]], on="package", how="left", suffixes=("","_price"))
    # statuses
    def status(val, p25, p50, p75):
        if pd.isna(val) or pd.isna(p50):
            return "n/a"
        if val > p75: return "higher than peers"
        if val < p25: return "lower than peers"
        return "about average"
    out["lead_status"] = out.apply(lambda r: status(r.get("lead_time_days"), r.get("lead_p25"), r.get("lead_p50"), r.get("lead_p75")), axis=1)
    if "price" in out.columns:
        out["price_status"] = out.apply(lambda r: status(r.get("price"), r.get("price_p25"), r.get("price_p50"), r.get("price_p75")), axis=1)
    # pct deltas
    def pct(val, p50):
        if pd.isna(val) or pd.isna(p50) or p50 == 0: return np.nan
        return (val - p50) / p50
    out["lead_pct_vs_p50"] = out.apply(lambda r: pct(r.get("lead_time_days"), r.get("lead_p50")), axis=1)
    if "price" in out.columns:
        out["price_pct_vs_p50"] = out.apply(lambda r: pct(r.get("price"), r.get("price_p50")), axis=1)
    return out

def vendor_mix_heatmap(proc_df: pd.DataFrame, peer_ids: set, my_proc: Optional[pd.DataFrame] = None):
    """
    Build a vendor x package matrix (counts) for peers; optional overlay for your vendors.
    """
    df = proc_df.copy()
    df = df[df["project_id"].astype(str).isin({str(x) for x in peer_ids})]
    if df.empty:
        return None, None
    pivot = df.pivot_table(index="vendor", columns="package", values="project_id", aggfunc="count", fill_value=0)
    overlay = None
    if my_proc is not None and not my_proc.empty and "vendor" in my_proc.columns and "package" in my_proc.columns:
        overlay = my_proc.pivot_table(index="vendor", columns="package", values="package", aggfunc="count", fill_value=0)
        # Align to same axes
        overlay = overlay.reindex(index=pivot.index, columns=pivot.columns, fill_value=0)
    return pivot, overlay

# -----------------------------
# UI
# -----------------------------
def render():
    st.title("🏷️ AI Benchmarking — Cost • Schedule • Economics • Procurement")

    # Project context for AI filter proposals
    pid = st.session_state.get("current_project_id") or "P-DEMO"
    ph  = st.session_state.get("current_phase_id") or f"PH-{st.session_state.get('fel_stage','FEL1')}"
    ctx_text = format_context_text(collect_phase_context(pid, ph))

    with st.expander("AI Peer-Set Proposal"):
        if st.button("Ask AI to propose filters", key=k("ai","ask")):
            st.session_state[k("ai","props")] = ai_propose_filters(ctx_text)
        props = st.session_state.get(k("ai","props"))
        if props:
            st.json(props)

    st.markdown("### 1) Upload datasets")
    c1, c2, c3 = st.columns(3)
    f_projects   = c1.file_uploader("projects.csv", type=["csv"], key=k("up","proj"))
    f_components = c2.file_uploader("components.csv", type=["csv"], key=k("up","comp"))
    f_proc       = c3.file_uploader("procurement.csv (optional peers)", type=["csv"], key=k("up","proc"))
    c4, c5, c6 = st.columns(3)
    f_my_comps  = c4.file_uploader("my_components.csv (your project)", type=["csv"], key=k("up","mine"))
    f_my_proc   = c5.file_uploader("my_procurement.csv (optional)", type=["csv"], key=k("up","myproc"))
    f_fx        = c6.file_uploader("fx.csv (currency→USD)", type=["csv"], key=k("up","fx"))

    # Read
    projs = pd.read_csv(f_projects) if f_projects else pd.DataFrame()
    comps = pd.read_csv(f_components) if f_components else pd.DataFrame()
    proc  = pd.read_csv(f_proc) if f_proc else pd.DataFrame()
    mine  = pd.read_csv(f_my_comps) if f_my_comps else pd.DataFrame()
    mypr  = pd.read_csv(f_my_proc) if f_my_proc else pd.DataFrame()
    fx    = pd.read_csv(f_fx) if f_fx else pd.DataFrame()

    # Normalize to USD if FX has been provided
    if not fx.empty:
        projs = normalize_currency(projs, fx, ["total_capex_usd"], currency_col="currency") if "total_capex_usd" in projs.columns else projs
        comps = normalize_currency(comps, fx, ["unit_cost","install_cost"], currency_col="currency")
        proc  = normalize_currency(proc,  fx, ["price"], currency_col="currency")
        mine  = normalize_currency(mine, fx, ["unit_cost","install_cost"], currency_col="currency")
        mypr  = normalize_currency(mypr, fx, ["price"], currency_col="currency")

    st.markdown("### 2) Select / confirm peer-set filters")
    props = st.session_state.get(k("ai","props")) or {}
    sector = st.text_input("Sector", value=props.get("sector",""), key=k("flt","sect"))
    region = st.text_input("Region", value=props.get("region",""), key=k("flt","reg"))
    complexity = st.selectbox("Complexity", ["low","medium","high","very high"],
                              index=["low","medium","high","very high"].index(props.get("complexity","medium") if props else "medium"),
                              key=k("flt","cx"))
    scope_tags = st.text_input("Scope tags (comma-separated)", value=props.get("scope_tags",""), key=k("flt","tags"))

    st.markdown("### 3) Compute similar projects")
    topn = st.slider("How many peers?", 3, 20, 10, key=k("peer","n"))
    if st.button("Find Similar Projects", use_container_width=True, key=k("peer","go")):
        if projs.empty:
            st.error("Upload projects.csv first.")
            return

        tgt_sector = sector.strip().lower()
        tgt_region = region.strip().lower()
        tgt_complexity = complexity.strip().lower()
        tgt_tags_set = parse_tags(scope_tags)

        df = projs.copy()
        df["_sim"] = df.apply(lambda r: similarity_row(r, tgt_sector, tgt_region, tgt_complexity, tgt_tags_set), axis=1)
        peers = df.sort_values("_sim", ascending=False).head(topn)
        st.session_state[k("peer","df")] = peers
        st.success(f"Found {len(peers)} similar projects.")
        st.dataframe(peers[["project_id","company","sector","region","complexity","scope_tags","total_capex_usd","total_duration_days","_sim"]], use_container_width=True)

        # Downloads
        buf = io.BytesIO(); peers.to_csv(buf, index=False)
        st.download_button("Download similar_projects.csv", data=buf.getvalue(), file_name="similar_projects.csv", mime="text/csv", key=k("peer","dl"))

        # Charts
        if PLOTLY_OK:
            if "total_capex_usd" in peers.columns:
                fig = px.box(peers, y="total_capex_usd", points="all", title="Peer CAPEX (USD)")
                st.plotly_chart(fig, use_container_width=True)
            if "total_duration_days" in peers.columns:
                fig2 = px.box(peers, y="total_duration_days", points="all", title="Peer Total Duration (days)")
                st.plotly_chart(fig2, use_container_width=True)

    st.markdown("### 4) Component benchmarking (pump, compressor, pipe, etc.)")
    if st.button("Compare Components", use_container_width=True, key=k("cmp","go")):
        peers = st.session_state.get(k("peer","df"))
        if peers is None or peers.empty:
            st.error("Run 'Find Similar Projects' first.")
            return
        if comps.empty:
            st.error("Upload components.csv for the benchmark set.")
            return

        peer_ids = set(peers["project_id"].astype(str))
        bench = comps[comps["project_id"].astype(str).isin(peer_ids)].copy()
        if bench.empty:
            st.warning("No components found for selected peers.")
        else:
            st.dataframe(bench.head(50), use_container_width=True)

            results = []
            if not mine.empty:
                my = mine.copy()
                for c in ["unit_cost","install_cost","duration_days"]:
                    if c in my.columns: my[c] = pd.to_numeric(my[c], errors="coerce")
                for c in ["unit_cost","install_cost","duration_days"]:
                    if c in bench.columns: bench[c] = pd.to_numeric(bench[c], errors="coerce")
                for _, r in my.iterrows():
                    ctype = str(r.get("component_type","")).strip().lower()
                    size  = str(r.get("spec_size","")).strip().lower()
                    subset = bench[bench["component_type"].str.strip().str.lower() == ctype].copy()
                    for metric in ["unit_cost","install_cost","duration_days"]:
                        if metric in r and metric in subset.columns:
                            comp = compare_value_to_peer(r.get(metric), subset[metric])
                            results.append({
                                "component_type": ctype,
                                "spec_size": size,
                                "metric": metric,
                                "my_value": r.get(metric),
                                "peer_p50": comp["p50"],
                                "peer_p25": comp["p25"],
                                "peer_p75": comp["p75"],
                                "peer_p90": comp.get("p90", np.nan),
                                "status": comp["status"],
                                "pct_vs_p50": comp["pct_vs_p50"],
                                "n_peers": int(subset[metric].dropna().shape[0]),
                            })
            res_df = pd.DataFrame(results)
            st.session_state[k("cmp","res")] = res_df
            if not res_df.empty:
                st.markdown("**Component Comparison Results**")
                st.dataframe(res_df, use_container_width=True)
                buf = io.BytesIO(); res_df.to_csv(buf, index=False)
                st.download_button("Download component_comparison.csv", data=buf.getvalue(),
                                   file_name="component_comparison.csv", mime="text/csv", key=k("cmp","dl"))

            # Optional visuals
            if PLOTLY_OK:
                for metric in ["unit_cost","install_cost","duration_days"]:
                    if metric in bench.columns:
                        try:
                            fig = px.box(bench, x="component_type", y=metric, points="outliers", title=f"Peers — {metric}")
                            st.plotly_chart(fig, use_container_width=True)
                        except Exception:
                            pass

    # ---------------- NEW: Procurement benchmarking ----------------
    st.markdown("### 4.1) Procurement benchmarking (lead-time & price)")
    if st.button("Analyze Procurement", use_container_width=True, key=k("proc","go")):
        peers = st.session_state.get(k("peer","df"))
        if peers is None or peers.empty:
            st.error("Run 'Find Similar Projects' first.")
            return
        if proc.empty:
            st.error("Upload procurement.csv to benchmark peers.")
            return

        peer_ids = set(peers["project_id"].astype(str))
        lead_stats, price_stats = procurement_peer_stats(proc, peer_ids)
        if lead_stats is None and price_stats is None:
            st.warning("No lead-time or price data available in procurement.csv.")
        else:
            if lead_stats is not None and not lead_stats.empty:
                st.markdown("**Peer lead-time by package (days)**")
                st.dataframe(lead_stats.sort_values("p50"), use_container_width=True)
                b1 = io.BytesIO(); lead_stats.to_csv(b1, index=False)
                st.download_button("Download peer_leadtime_stats.csv", data=b1.getvalue(),
                                   file_name="peer_leadtime_stats.csv", mime="text/csv", key=k("proc","dl_lead"))
                if PLOTLY_OK:
                    try:
                        f = px.bar(lead_stats.sort_values("p50"), x="package", y="p50", title="Peer median lead-time (days)")
                        st.plotly_chart(f, use_container_width=True)
                    except Exception:
                        pass

            if price_stats is not None and not price_stats.empty:
                st.markdown("**Peer price by package (USD)**")
                st.dataframe(price_stats.sort_values("p50"), use_container_width=True)
                b2 = io.BytesIO(); price_stats.to_csv(b2, index=False)
                st.download_button("Download peer_price_stats.csv", data=b2.getvalue(),
                                   file_name="peer_price_stats.csv", mime="text/csv", key=k("proc","dl_price"))
                if PLOTLY_OK:
                    try:
                        f2 = px.bar(price_stats.sort_values("p50"), x="package", y="p50", title="Peer median price (USD)")
                        st.plotly_chart(f2, use_container_width=True)
                    except Exception:
                        pass

            # Compare your procurement (if provided)
            if not mypr.empty:
                my_vs = compare_my_procurement(mypr, lead_stats, price_stats)
                st.markdown("**Your procurement vs peer stats**")
                st.dataframe(my_vs, use_container_width=True)
                b3 = io.BytesIO(); my_vs.to_csv(b3, index=False)
                st.download_button("Download my_vs_peers_procurement.csv", data=b3.getvalue(),
                                   file_name="my_vs_peers_procurement.csv", mime="text/csv", key=k("proc","dl_mine"))

    st.markdown("### 4.2) Vendor mix heatmap")
    if st.button("Build Vendor Heatmap", use_container_width=True, key=k("vend","go")):
        peers = st.session_state.get(k("peer","df"))
        if peers is None or peers.empty:
            st.error("Run 'Find Similar Projects' first.")
            return
        if proc.empty:
            st.error("Upload procurement.csv to build vendor heatmap.")
            return
        peer_ids = set(peers["project_id"].astype(str))
        pivot, overlay = vendor_mix_heatmap(proc, peer_ids, mypr if not mypr.empty else None)
        if pivot is None or pivot.empty:
            st.warning("Not enough peer vendor/package data to build heatmap.")
        else:
            st.markdown("**Peer vendor mix (count of packages per vendor)**")
            st.dataframe(pivot, use_container_width=True)
            if PLOTLY_OK:
                try:
                    fig = go.Figure(data=go.Heatmap(
                        z=pivot.values, x=list(pivot.columns), y=list(pivot.index),
                        colorbar_title="count"
                    ))
                    fig.update_layout(title="Vendor × Package (Peers)")
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    pass

            if overlay is not None and not overlay.empty and PLOTLY_OK:
                try:
                    # Overlay markers where you have vendor-package pairs
                    Y, X = np.where(overlay.values > 0)
                    overlay_points = pd.DataFrame({
                        "vendor": [overlay.index[i] for i in Y],
                        "package": [overlay.columns[j] for j in X],
                        "count": [overlay.values[i, j] for i, j in zip(Y, X)]
                    })
                    st.markdown("**Your vendor-package usage (overlay markers)**")
                    st.dataframe(overlay_points, use_container_width=True)
                    # Add as scattergl over heatmap axes
                    fig.add_trace(go.Scatter(
                        x=overlay_points["package"], y=overlay_points["vendor"],
                        mode="markers", name="your usage", marker=dict(size=10, symbol="x")
                    ))
                    st.plotly_chart(fig, use_container_width=True)
                except Exception:
                    pass

            # Downloads
            h1 = io.BytesIO(); pivot.to_csv(h1)
            st.download_button("Download vendor_mix_peers.csv", data=h1.getvalue(),
                               file_name="vendor_mix_peers.csv", mime="text/csv", key=k("vend","dl_peers"))
            if overlay is not None and not overlay.empty:
                h2 = io.BytesIO(); overlay.to_csv(h2)
                st.download_button("Download vendor_mix_mine.csv", data=h2.getvalue(),
                                   file_name="vendor_mix_mine.csv", mime="text/csv", key=k("vend","dl_mine"))

    # ---------------- AI Summary ----------------
    st.markdown("### 5) AI Review Summary & Recommendations")
    if st.button("Generate AI Benchmarking Summary", use_container_width=True, key=k("ai","sum")):
        peers = st.session_state.get(k("peer","df"))
        cmpres = st.session_state.get(k("cmp","res"))
        llm = get_chat_callable()
        system = (
            "You are a benchmarking analyst. Write a concise PMO review (<= 250 words) covering: "
            "peer selection rationale; cost & schedule vs peers; component outliers (pump/compressor/piping); "
            "procurement lead-times and prices vs peers; vendor mix insights; "
            "and 5 prioritized actions to reduce cost & duration. Use plain language."
        )
        payload = {
            "filters": {"sector": sector, "region": region, "complexity": complexity, "scope_tags": scope_tags},
            "peers": (peers.head(12).to_dict(orient="records") if peers is not None and not peers.empty else []),
            "components_review": (cmpres.to_dict(orient="records") if cmpres is not None and not cmpres.empty else []),
        }
        # Add procurement snippets if available
        try:
            lead_stats = st.session_state.get(k("cache","lead_stats"))
        except Exception:
            lead_stats = None
        try:
            price_stats = st.session_state.get(k("cache","price_stats"))
        except Exception:
            price_stats = None

        # Grab live from screen if needed
        # (not caching inside session; OK to skip if not present)

        user = f"PROJECT CONTEXT:\n{ctx_text[:4000]}\n\nDERIVED DATA (compact):\n{json.dumps(payload, indent=2)[:8000]}"
        try:
            text = str(llm(user, system))
        except Exception as e:
            text = f"(AI review unavailable: {e})"
        st.text_area("AI Benchmarking Summary (editable)", value=text, height=260, key=k("ai","text"))

        if DOCX_OK:
            b = _docx_bytes("AI Benchmarking Review", st.session_state.get(k("ai","text")) or text)
            if b:
                st.download_button("Download DOCX", data=b,
                                   file_name="ai_benchmarking_review.docx",
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key=k("ai","docx"))

        _save_artifact_button(
            "AI_Benchmarking_Review",
            {"filters": payload["filters"],
             "peers": payload["peers"],
             "components_review": payload["components_review"],
             "summary": st.session_state.get(k("ai","text")) or text},
            pid, ph, key=k("ai","save")
        )

    st.markdown("---")
    st.markdown("#### CSV schema cheat-sheet")
    st.code(json.dumps({
        "projects.csv": [
            "project_id","company","sector","region","complexity","scope_tags",
            "total_capex_usd","total_duration_days","start_year","currency?","source_url?"
        ],
        "components.csv": [
            "project_id","component_type","spec_size","count","unit_cost","install_cost",
            "duration_days","vendor?","currency?"
        ],
        "procurement.csv (peers)": [
            "project_id","package","vendor","lead_time_days?","price?","currency?"
        ],
        "my_components.csv (your project)": [
            "component_type","spec_size","count","unit_cost","install_cost","duration_days","vendor?","currency?"
        ],
        "my_procurement.csv (your project)": [
            "package","vendor","lead_time_days?","price?","currency?"
        ],
        "fx.csv (optional)": ["currency","to_usd_rate"]
    }, indent=2), language="json")
