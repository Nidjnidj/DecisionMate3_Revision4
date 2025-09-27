# ai_pm_analytics.py
"""
DecisionMate — AI PM Analytics
Data-driven PM insights: EVM (SPI/CPI), schedule & cost variance, trends, burndown,
anomaly detection, and Monte Carlo schedule risk — plus AI-generated PM narrative.

Inputs (upload any subset; the more the better):
  progress.csv      : date, planned_value, earned_value, actual_cost
  tasks.csv         : task_id, name, start (YYYY-MM-DD), finish (YYYY-MM-DD),
                      duration_days, percent_complete, optimistic_days?, most_likely_days?, pessimistic_days?
  risks.csv         : risk_id, description, probability (0-1), impact_cost, impact_days, owner, status
  issues.csv        : issue_id, title, severity, opened_on, closed_on, owner, status
  timesheets.csv    : date, person, task_id, hours
  costs.csv         : date, category, amount, type (CAPEX/OPEX)

Install (if needed):
    pip install plotly
"""

from __future__ import annotations
import io, uuid, math, json, random
from dataclasses import dataclass
from typing import Optional, Dict, List, Tuple

import streamlit as st
import pandas as pd
import numpy as np

# --- Optional DOCX export ---
try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# --- App services (LLM + context) ---
from .providers import get_chat_callable
from .tools import collect_phase_context, format_context_text

# -----------------------------
# Namespaced keys (avoid collisions)
# -----------------------------
if "AI_PM_NS" not in st.session_state:
    st.session_state["AI_PM_NS"] = f"pma_{uuid.uuid4().hex[:8]}"

def k(group: str, name: str) -> str:
    return f"{st.session_state['AI_PM_NS']}::{group}::{name}"

# -----------------------------
# Helpers
# -----------------------------
def _docx_bytes(title: str, body_markdown: str) -> Optional[bytes]:
    if not DOCX_OK:
        return None
    try:
        doc = Document()
        doc.add_heading(title or "PM Analytics Report", level=1)
        for line in (body_markdown or "").splitlines():
            if line.strip().startswith("#"):
                doc.add_heading(line.strip("# ").strip(), level=2)
            else:
                doc.add_paragraph(line)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None

def _save_artifact_button(artifact_type: str, payload: dict, pid: str, ph: str, key: str):
    if st.button("Save as Artifact", key=key):
        try:
            from artifact_registry import save
            save(
                project_id=pid,
                phase_id=ph,
                artifact_type=artifact_type,
                data=payload,
                status="Draft",
            )
            st.success(f"Saved artifact: {artifact_type} (Draft)")
        except Exception as e:
            st.error(f"Could not save artifact: {e}")

# -----------------------------
# EVM / metrics
# -----------------------------
def compute_evm(progress_df: pd.DataFrame) -> pd.DataFrame:
    """Return EVM metrics per date with SPI, CPI, SV, CV, EAC, ETC, TCPI."""
    df = progress_df.copy()
    req = {"date","planned_value","earned_value","actual_cost"}
    if not req.issubset(set(df.columns)):
        return pd.DataFrame()
    df["date"] = pd.to_datetime(df["date"])
    df = df.sort_values("date")
    df["PV"] = df["planned_value"].astype(float)
    df["EV"] = df["earned_value"].astype(float)
    df["AC"] = df["actual_cost"].astype(float)
    df["SV"] = df["EV"] - df["PV"]                      # schedule variance
    df["CV"] = df["EV"] - df["AC"]                      # cost variance
    df["SPI"] = df["EV"] / df["PV"].replace({0: np.nan})
    df["CPI"] = df["EV"] / df["AC"].replace({0: np.nan})
    BAC = df["PV"].max() if len(df) else 0.0            # budget at completion (from PV curve)
    df["EAC"] = np.where(df["CPI"] > 0, BAC / df["CPI"], np.nan)
    df["ETC"] = df["EAC"] - df["AC"]
    df["TCPI"] = (BAC - df["EV"]) / (BAC - df["AC"]).replace({0: np.nan})
    return df

def anomaly_flags(series: pd.Series, z_thresh: float = 2.5) -> pd.Series:
    """Return boolean flags where abs(zscore) > threshold."""
    s = series.astype(float)
    mu, sd = s.mean(), s.std(ddof=0)
    if sd == 0 or np.isnan(sd):
        return pd.Series([False]*len(s), index=s.index)
    z = (s - mu) / sd
    return z.abs() > z_thresh

# -----------------------------
# Critical path (simple)
# -----------------------------
def critical_path(tasks_df: pd.DataFrame) -> Tuple[pd.DataFrame, List[str]]:
    """
    Simple longest-path by finish_day over DAG from predecessor_ids.
    Expects: task_id, duration_days, predecessor_ids (comma-separated)
    Returns: tasks with ES/EF/LS/LF and critical bool; and the critical chain IDs.
    """
    if tasks_df.empty or "task_id" not in tasks_df.columns:
        return pd.DataFrame(), []
    df = tasks_df.copy()
    df["duration_days"] = df["duration_days"].fillna(0).astype(int)
    # parse predecessors
    preds = {}
    for _, r in df.iterrows():
        s = str(r.get("predecessor_ids", "") or "").strip()
        preds[r["task_id"]] = [p.strip() for p in s.split(",") if p.strip()]
    # Topo order (Kahn)
    nodes = set(df["task_id"].tolist())
    indeg = {n: 0 for n in nodes}
    for n, ps in preds.items():
        for p in ps:
            if p in nodes:
                indeg[n] += 1
    Q = [n for n in nodes if indeg[n] == 0]
    order = []
    while Q:
        x = Q.pop(0)
        order.append(x)
        for n, ps in preds.items():
            if x in ps:
                indeg[n] -= 1
                if indeg[n] == 0:
                    Q.append(n)
    if len(order) != len(nodes):
        # cycle detected or missing predecessors; bail gracefully
        return pd.DataFrame(), []

    # Forward pass (ES/EF)
    ES, EF = {}, {}
    for t_id in order:
        dur = int(df.loc[df["task_id"]==t_id, "duration_days"].iloc[0] or 0)
        es = 0
        for p in preds[t_id]:
            es = max(es, EF.get(p, 0))
        ES[t_id] = es
        EF[t_id] = es + dur

    # Backward pass (LS/LF)
    proj_finish = max(EF.values()) if EF else 0
    LS, LF = {}, {}
    for t_id in reversed(order):
        dur = int(df.loc[df["task_id"]==t_id, "duration_days"].iloc[0] or 0)
        # if no successors, LF = proj_finish
        succs = [n for n, ps in preds.items() if t_id in ps]
        if not succs:
            LF[t_id] = proj_finish
        else:
            LF[t_id] = min(LS.get(s, proj_finish) for s in succs)
        LS[t_id] = LF[t_id] - dur

    out = df[["task_id","name","duration_days"]].copy()
    out["ES"] = out["task_id"].map(ES)
    out["EF"] = out["task_id"].map(EF)
    out["LS"] = out["task_id"].map(LS)
    out["LF"] = out["task_id"].map(LF)
    out["TotalFloat"] = out["LS"] - out["ES"]
    out["is_critical"] = out["TotalFloat"] <= 0
    crit_chain = []
    # reconstruct one critical chain (greedy by EF)
    cur = None
    # pick task with EF == proj_finish and is critical
    cand = out[out["is_critical"] & (out["EF"] == proj_finish)]
    if not cand.empty:
        cur = cand.iloc[0]["task_id"]
        crit_chain.append(cur)
        while True:
            prevs = preds[cur]
            prev_crit = out[out["task_id"].isin(prevs) & out["is_critical"]]
            if prev_crit.empty:
                break
            nxt = prev_crit.sort_values("EF", ascending=False).iloc[0]["task_id"]
            crit_chain.append(nxt)
            cur = nxt
        crit_chain.reverse()
    return out.sort_values("ES"), crit_chain

# -----------------------------
# Monte Carlo schedule risk
# -----------------------------
def triangular_sample(a, m, b):
    return np.random.triangular(a, m, b)

def simulate_schedule(tasks_df: pd.DataFrame, iters: int = 1000) -> pd.DataFrame:
    """
    PERT/triangular simulation over durations.
    If optimistic/most_likely/pessimistic not provided, fall back to +/- 20% around duration_days.
    Returns a dataframe with simulated project finish (days) per iteration.
    """
    if tasks_df.empty:
        return pd.DataFrame()
    df = tasks_df.copy()
    df["duration_days"] = df["duration_days"].fillna(0).astype(float)

    # Parse predecessors
    preds = {}
    for _, r in df.iterrows():
        s = str(r.get("predecessor_ids", "") or "").strip()
        preds[r["task_id"]] = [p.strip() for p in s.split(",") if p.strip()]

    finishes = []
    for _ in range(iters):
        # sample durations
        dur = {}
        for _, r in df.iterrows():
            base = float(r["duration_days"])
            a = float(r.get("optimistic_days", np.nan))
            m = float(r.get("most_likely_days", np.nan))
            b = float(r.get("pessimistic_days", np.nan))
            if np.isnan(a) or np.isnan(m) or np.isnan(b) or a <= 0 or b <= 0:
                # fallback +/-20%
                a = max(1.0, 0.8 * base)
                m = max(1.0, base)
                b = max(1.0, 1.2 * base)
            dur[r["task_id"]] = triangular_sample(a, m, b)

        # topo + forward pass to compute finish
        nodes = set(df["task_id"].tolist())
        indeg = {n: 0 for n in nodes}
        for n, ps in preds.items():
            for p in ps:
                if p in nodes:
                    indeg[n] += 1
        Q = [n for n in nodes if indeg[n]==0]
        order = []
        while Q:
            x = Q.pop(0)
            order.append(x)
            for n, ps in preds.items():
                if x in ps:
                    indeg[n] -= 1
                    if indeg[n]==0:
                        Q.append(n)
        ES, EF = {}, {}
        for t_id in order:
            es = 0.0
            for p in preds[t_id]:
                es = max(es, EF.get(p, 0.0))
            ES[t_id] = es
            EF[t_id] = es + float(dur[t_id])
        finishes.append(max(EF.values()) if EF else 0.0)

    return pd.DataFrame({"iteration": np.arange(1, len(finishes)+1), "finish_days": finishes})

# -----------------------------
# AI PM narrative
# -----------------------------
def ai_pm_narrative(context_text: str, evm_tail: pd.Series, flags: Dict[str, bool], risks_top: pd.DataFrame) -> str:
    llm = get_chat_callable()
    system = (
        "You are a PMO advisor. Write a crisp weekly PM status narrative (<= 200 words): "
        "start with health (SPI/CPI), then schedule vs plan, cost vs budget, top risks/issues, "
        "and 3-5 actionable next steps. Plain text; no markdown; no tables."
    )
    tail = evm_tail.fillna("").to_dict() if evm_tail is not None else {}
    user = f"""
PROJECT CONTEXT:
{context_text[:5000]}

LATEST EVM SNAPSHOT:
{json.dumps({k: (float(v) if isinstance(v,(int,float,np.floating)) else str(v)) for k,v in tail.items()}, indent=2)}

ANOMALY FLAGS:
{json.dumps(flags, indent=2)}

TOP RISKS (if any):
{risks_top.to_dict(orient='records') if risks_top is not None else []}
"""
    try:
        return str(llm(user, system))
    except Exception as e:
        return f"(AI narrative unavailable: {e})"

# -----------------------------
# UI
# -----------------------------
def render():
    st.title("📊 AI PM Analytics — EVM • Trends • Risk • Monte Carlo • Narrative")

    # Project context for AI narrative
    pid = st.session_state.get("current_project_id") or "P-DEMO"
    ph = st.session_state.get("current_phase_id") or f"PH-{st.session_state.get('fel_stage','FEL1')}"
    ctx_text = format_context_text(collect_phase_context(pid, ph))

    st.markdown("### 1) Upload PM data (any subset)")
    c1, c2, c3 = st.columns(3)
    f_progress = c1.file_uploader("progress.csv", type=["csv"], key=k("up","progress"))
    f_tasks    = c2.file_uploader("tasks.csv", type=["csv"], key=k("up","tasks"))
    f_costs    = c3.file_uploader("costs.csv (optional)", type=["csv"], key=k("up","costs"))
    c4, c5, c6 = st.columns(3)
    f_times    = c4.file_uploader("timesheets.csv (optional)", type=["csv"], key=k("up","times"))
    f_risks    = c5.file_uploader("risks.csv (optional)", type=["csv"], key=k("up","risks"))
    f_issues   = c6.file_uploader("issues.csv (optional)", type=["csv"], key=k("up","issues"))

    # Read data
    progress = pd.read_csv(f_progress) if f_progress else pd.DataFrame()
    tasks    = pd.read_csv(f_tasks) if f_tasks else pd.DataFrame()
    costs    = pd.read_csv(f_costs) if f_costs else pd.DataFrame()
    times    = pd.read_csv(f_times) if f_times else pd.DataFrame()
    risks    = pd.read_csv(f_risks) if f_risks else pd.DataFrame()
    issues   = pd.read_csv(f_issues) if f_issues else pd.DataFrame()

    st.markdown("### 2) Core metrics & trends (EVM)")
    evm = compute_evm(progress) if not progress.empty else pd.DataFrame()
    if evm.empty:
        st.info("Upload progress.csv with columns: date, planned_value, earned_value, actual_cost")
    else:
        st.dataframe(evm.tail(12), use_container_width=True)

        # Plotly S-curve & indices
        try:
            import plotly.express as px
            # S-curve
            sfig = px.line(evm, x="date", y=["PV","EV","AC"], markers=True, title="S-Curve (PV/EV/AC)")
            st.plotly_chart(sfig, use_container_width=True)
            # SPI/CPI
            icols = ["SPI","CPI"]
            if evm[icols].notna().any().any():
                ifig = px.line(evm, x="date", y=icols, markers=True, title="SPI & CPI over time")
                st.plotly_chart(ifig, use_container_width=True)
        except Exception:
            pass

        # Latest snapshot & anomaly flags
        tail = evm.iloc[-1]
        flags = {
            "SPI_below_0.9": bool(tail.get("SPI", 1.0) < 0.9),
            "CPI_below_0.9": bool(tail.get("CPI", 1.0) < 0.9),
            "SV_negative":   bool(tail.get("SV", 0.0) < 0),
            "CV_negative":   bool(tail.get("CV", 0.0) < 0),
        }
        st.markdown(f"**Latest:** Date={tail['date'].date() if 'date' in tail else ''} | "
                    f"SPI={tail.get('SPI',np.nan):.2f} | CPI={tail.get('CPI',np.nan):.2f} | "
                    f"SV={tail.get('SV',np.nan):,.0f} | CV={tail.get('CV',np.nan):,.0f} | "
                    f"EAC={tail.get('EAC',np.nan):,.0f}")

        # Anomaly detection on AC deltas (simple z-score)
        evm["AC_delta"] = evm["AC"].diff()
        evm["AC_anomaly"] = anomaly_flags(evm["AC_delta"].fillna(0.0))
        if evm["AC_anomaly"].any():
            st.warning("Cost anomalies detected in periods: " +
                       ", ".join(evm.loc[evm["AC_anomaly"], "date"].dt.date.astype(str)))

        # Downloads
        buf = io.BytesIO(); evm.to_csv(buf, index=False)
        st.download_button("Download evm_metrics.csv", data=buf.getvalue(),
                           file_name="evm_metrics.csv", mime="text/csv", key=k("dl","evm"))

    st.markdown("### 3) Schedule: critical path (if tasks uploaded)")
    if not tasks.empty:
        cp_df, chain = critical_path(tasks)
        if cp_df.empty:
            st.info("For critical path, include columns: task_id, name, duration_days, predecessor_ids")
        else:
            st.dataframe(cp_df, use_container_width=True)
            if chain:
                st.success("Critical chain: " + " → ".join(chain))

            # Burndown (by % complete if provided)
            if "percent_complete" in tasks.columns:
                try:
                    import plotly.express as px
                    tcopy = tasks.copy()
                    tcopy["percent_complete"] = pd.to_numeric(tcopy["percent_complete"], errors="coerce").fillna(0.0)
                    tcopy["remaining"] = 100.0 - tcopy["percent_complete"]
                    bfig = px.bar(tcopy.sort_values("remaining", ascending=False),
                                  x="name", y="remaining", title="Remaining Work by Task (%)")
                    st.plotly_chart(bfig, use_container_width=True)
                except Exception:
                    pass

    st.markdown("### 4) Risk & issues (if uploaded)")
    top_risks = pd.DataFrame()
    if not risks.empty:
        r = risks.copy()
        # crude risk exposure
        if {"probability","impact_cost"}.issubset(set(r.columns)):
            r["exposure_cost"] = r["probability"].astype(float) * r["impact_cost"].astype(float)
        if {"probability","impact_days"}.issubset(set(r.columns)):
            r["exposure_days"] = r["probability"].astype(float) * r["impact_days"].astype(float)
        top_risks = r.sort_values(["exposure_cost","exposure_days"], ascending=False, na_position="last").head(10)
        st.dataframe(top_risks, use_container_width=True)
    if not issues.empty:
        st.dataframe(issues.sort_values("severity", ascending=False), use_container_width=True)

    st.markdown("### 5) Monte Carlo schedule risk (optional)")
    iters = st.slider("Iterations", 200, 3000, 1000, 100, key=k("mc","iters"))
    if not tasks.empty and st.button("Run Simulation", key=k("mc","run"), use_container_width=True):
        sims = simulate_schedule(tasks, iters=iters)
        if sims.empty:
            st.info("Provide tasks with duration and predecessors to simulate.")
        else:
            p50 = np.percentile(sims["finish_days"], 50)
            p80 = np.percentile(sims["finish_days"], 80)
            p90 = np.percentile(sims["finish_days"], 90)
            st.write(f"**Finish (days):** P50={p50:.0f}, P80={p80:.0f}, P90={p90:.0f}")

            try:
                import plotly.express as px
                hfig = px.histogram(sims, x="finish_days", nbins=30, title="Completion Distribution (days)")
                st.plotly_chart(hfig, use_container_width=True)
            except Exception:
                pass

            buf = io.BytesIO(); sims.to_csv(buf, index=False)
            st.download_button("Download simulation.csv", data=buf.getvalue(),
                               file_name="schedule_simulation.csv", mime="text/csv", key=k("dl","sim"))

    st.markdown("### 6) AI PM Narrative & Next Actions")
    if not evm.empty:
        tail = evm.iloc[-1]
    else:
        tail = pd.Series(dtype=object)
    flags = {"SPI<0.9": False, "CPI<0.9": False}
    if not evm.empty:
        flags["SPI<0.9"] = bool(tail.get("SPI", 1.0) < 0.9)
        flags["CPI<0.9"] = bool(tail.get("CPI", 1.0) < 0.9)
    narrative = ai_pm_narrative(ctx_text, tail, flags, top_risks if not top_risks.empty else None)
    st.text_area("AI Narrative (editable)", value=narrative, height=220, key=k("ai","narr"))

    # Exports & artifacts
    cols = st.columns(3)
    if DOCX_OK:
        data = _docx_bytes("PM Weekly Status", st.session_state.get(k("ai","narr")) or narrative)
        if data:
            cols[0].download_button("Download DOCX", data=data,
                                    file_name="pm_weekly_status.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=k("dl","docx"))
    # Save artifacts
    payload = {
        "evm": evm.to_dict(orient="records") if not evm.empty else [],
        "critical_path": top_risks.to_dict(orient="records") if not top_risks.empty else [],
        "narrative": st.session_state.get(k("ai","narr")) or narrative
    }
    _save_artifact_button("AI_PM_Analytics", payload, pid, ph, key=k("save","art"))

    st.markdown("---")
    st.markdown("#### Expected CSV schemas (minimal)")
    st.code(json.dumps({
        "progress.csv": ["date","planned_value","earned_value","actual_cost"],
        "tasks.csv": ["task_id","name","start","finish","duration_days","percent_complete",
                      "optimistic_days?","most_likely_days?","pessimistic_days?","predecessor_ids?"],
        "risks.csv": ["risk_id","description","probability","impact_cost","impact_days","owner","status"],
        "issues.csv": ["issue_id","title","severity","opened_on","closed_on","owner","status"],
        "timesheets.csv": ["date","person","task_id","hours"],
        "costs.csv": ["date","category","amount","type"]
    }, indent=2), language="json")
