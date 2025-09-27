import io
import uuid
from datetime import datetime
from typing import Optional
from . import ai_optimizer
from . import ai_benchmarking
import streamlit as st
from . import ai_pm_analytics
# ---- Core services / utils ----
from .providers import get_chat_callable
from .tools import collect_phase_context, format_context_text, read_uploaded_files
from .prompts import build_business_case_prompt

# ---- Optional: DOCX export ----
try:
    from docx import Document  # pip install python-docx
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# ---- Optional: RAG (only if files exist) ----
try:
    from .rag_indexer import rebuild_index
    from .rag_search import retrieve
    RAG_OK = True
except Exception:
    RAG_OK = False


# =========================
# Key namespacing & single-render guard
# =========================
if "AI_HUB_KEY_NS" not in st.session_state:
    st.session_state["AI_HUB_KEY_NS"] = f"hub_{uuid.uuid4().hex[:8]}"

def k(group: str, name: str) -> str:
    """Generate a unique, readable Streamlit key scoped to this page."""
    return f"{st.session_state['AI_HUB_KEY_NS']}::{group}::{name}"

# Prevent accidental double-render of this page within a single run
if st.session_state.get("ai_hub_rendered"):
    st.stop()
st.session_state["ai_hub_rendered"] = True


# =========================
# Helpers
# =========================
def _switch_tab(label: str):
    """Programmatically click a tab by its label after rerun."""
    js = f"""
    <script>
      const wanted = {label!r};
      const tabs = window.parent.document.querySelectorAll('button[role="tab"]');
      for (const t of tabs) {{
        if ((t.innerText || t.textContent).trim() === wanted) {{
          t.click();
          break;
        }}
      }}
    </script>"""
    st.components.v1.html(js, height=0)

def _active_ids():
    pid = st.session_state.get("current_project_id") or "P-DEMO"
    ph = st.session_state.get("current_phase_id") or f"PH-{st.session_state.get('fel_stage', 'FEL1')}"
    return pid, ph

@st.cache_data(ttl=60)
def _cached_phase_context(pid: str, ph: str) -> str:
    ctx = collect_phase_context(pid, ph)
    return format_context_text(ctx)

def _save_artifact_button(
    artifact_type: str,
    payload: dict,
    pid: str,
    ph: str,
    label: Optional[str] = None,
    key: Optional[str] = None
):
    label = label or "Save as Artifact"
    key = key or k("buttons", f"save_{artifact_type.lower()}")
    if st.button(label, key=key):
        try:
            from artifact_registry import save  # provided by your app
            save(
                project_id=pid,
                artifact_type=artifact_type,
                phase_id=ph,
                data=payload,
                status="Draft",
            )
            st.success(f"Saved as artifact: {artifact_type} (Draft)")
        except Exception as e:
            st.error(f"Could not save artifact: {e}")

def _docx_bytes(title: str, body_markdown: str) -> Optional[bytes]:
    """Create a DOCX file in-memory from markdown-like text."""
    if not DOCX_OK:
        return None
    try:
        doc = Document()
        doc.add_heading(title or "Document", level=1)
        for line in (body_markdown or "").splitlines():
            if line.strip().startswith("#"):
                doc.add_heading(line.strip("# ").strip(), level=2)
            else:
                doc.add_paragraph(line)
        doc.add_paragraph(f"\nGenerated: {datetime.now().isoformat(timespec='minutes')}")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except Exception:
        return None


# =========================
# Main page
# =========================
def render():
    st.title("🤖 DecisionMate — AI Services")
    st.caption("Context-aware assistant for your active project/phase.")

    pid, ph = _active_ids()
    llm = get_chat_callable()

    # ---- Sidebar ----
    with st.sidebar:
        st.markdown("### Context")
        st.write(f"**Project ID**: {pid}")
        st.write(f"**Phase ID**: {ph}")
        if st.button("Refresh context", key=k("sidebar", "refresh_ctx")):
            st.cache_data.clear()
            st.rerun()

        st.markdown("### AI Settings")
        temp = st.slider("Creativity", 0.0, 1.0, 0.2, 0.05, key=k("sidebar", "temp"))
        if not st.secrets.get("llm"):
            st.info("Configure in `.streamlit/secrets.toml` under `[llm]`.")

        if RAG_OK:
            st.markdown("### Knowledge Index (RAG)")
            st.caption("Optional: builds a local searchable index from your artifacts.")
            if st.button("Rebuild Knowledge Index", key=k("sidebar", "rag_rebuild")):
                msg = rebuild_index(pid, ph)
                st.info(msg)

    # ---- Context (cached) ----
    ctx_text = _cached_phase_context(pid, ph)

    # ---- Tabs ----
    TAB_LABELS = [
        "Chat with your Project",
        "Generate Executive Summary",
        "Risk & Next Actions",
        "Business Case Draft",
        "Guided Business Case (Q&A)",
        "AI Optimizer",
        "AI PM Analytics",
        "AI Benchmarking", 
    ]
    tabs = st.tabs(TAB_LABELS)

    # Re-select tab requested before rerun (if any)
    _force = st.session_state.pop("force_tab", None)
    if _force:
        # One-time restore per rerun
        if not st.session_state.get(k("tabs", "restored")):
            _switch_tab(_force)
            st.session_state[k("tabs", "restored")] = True

    # --------------------------
    # 1) Chat
    # --------------------------
    with tabs[0]:
        st.subheader("Chat with your Project")
        q = st.text_area(
            "Ask anything (the AI sees your current artifacts):",
            height=120,
            placeholder="e.g., What are we missing for gate readiness?",
            key=k("chat", "q"),
        )
        if st.button("Ask AI", use_container_width=True, key=k("chat", "ask")):
            system = (
                "You are DecisionMate AI. Be concise, practical, and action-oriented. "
                f"Creativity level: {st.session_state.get(k('sidebar','temp'), 0.2)}. Use the project context to answer accurately."
            )
            prompt = f"PROJECT CONTEXT:\n{ctx_text}\n\nQUESTION:\n{q}"
            ans = llm(prompt, system)
            st.write(ans)
            _save_artifact_button(
                "AI_Answer",
                {"question": q, "answer": ans},
                pid,
                ph,
                label="Save Answer as Artifact",
                key=k("chat", "save_answer_art"),
            )

    # --------------------------
    # 2) Executive Summary
    # --------------------------
    with tabs[1]:
        st.subheader("Executive Summary")
        bullets = st.checkbox("Use bullet style", value=True, key=k("sum", "bullets"))
        include_sources = st.checkbox("Include source hints (artifact types)", value=False, key=k("sum", "sources"))

        if st.button("Generate Summary", use_container_width=True, key=k("sum", "gen")):
            system = f"Create a crisp executive summary for a program review. Creativity: {st.session_state.get(k('sidebar','temp'), 0.2)}."
            prompt = (
                f"Context below. Produce a {'bullet-point' if bullets else 'narrative'} summary."
                + (" Add brief source hints like [Artifact:Type]." if include_sources else "")
                + f"\n\n{ctx_text}"
            )
            summary = llm(prompt, system)
            st.write(summary)

            _save_artifact_button(
                "AI_Exec_Summary",
                {"text": summary, "bullets": bullets, "include_sources": include_sources},
                pid,
                ph,
                key=k("sum", "save_art"),
            )

            if DOCX_OK:
                data = _docx_bytes("Executive Summary", summary)
                if data:
                    st.download_button(
                        "Download DOCX",
                        data=data,
                        file_name="executive_summary.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=k("sum", "dl"),
                    )
                else:
                    st.info("DOCX export not available on this system.")

    # --------------------------
    # 3) Risk & Next Actions
    # --------------------------
    with tabs[2]:
        st.subheader("Risk & Next Actions")
        gate = st.text_input(
            "Target gate/stage (optional)",
            value=st.session_state.get("fel_stage", "FEL1"),
            key=k("risk", "gate"),
        )
        n_risks = st.slider("How many risks?", 3, 10, 5, key=k("risk", "count"))
        if st.button("Analyze Risks and Suggest Next Actions", use_container_width=True, key=k("risk", "gen")):
            system = f"You are a PMO risk coach. Provide actionable risks and next steps. Creativity: {st.session_state.get(k('sidebar','temp'), 0.2)}."
            prompt = (
                f"Gate: {gate}\n\nPROJECT CONTEXT:\n{ctx_text}\n\n"
                f"Requirements:\n- List top {n_risks} risks with owner and due date.\n"
                "- Add 1–2 concrete next actions per risk.\n- If data is missing, mark items with ✅TODO."
            )
            risks = llm(prompt, system)
            st.write(risks)

            _save_artifact_button(
                "AI_Risk_Next_Actions",
                {"gate": gate, "text": risks, "n_risks": n_risks},
                pid,
                ph,
                key=k("risk", "save_art"),
            )

            if DOCX_OK:
                data = _docx_bytes("Risks & Next Actions", risks)
                if data:
                    st.download_button(
                        "Download DOCX",
                        data=data,
                        file_name="risks_next_actions.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=k("risk", "dl"),
                    )
                else:
                    st.info("DOCX export not available on this system.")

    # --------------------------
    # 4) Business Case Draft
    # --------------------------
    with tabs[3]:
        st.subheader("Business Case / Project Draft")

        industry = st.session_state.get("industry") or "general"
        default_gate = st.session_state.get("fel_stage", "FEL1")

        col1, col2 = st.columns(2)
        with col1:
            bc_title = st.text_input("Project Title", value=f"Project {pid}", key=k("bc", "title"))
            audience = st.text_input("Audience", value="Executive Steering Committee", key=k("bc", "audience"))
        with col2:
            target_gate = st.text_input("Target Gate/Stage", value=default_gate, key=k("bc", "gate"))
            tone = st.selectbox("Tone", ["concise, executive", "neutral, formal", "persuasive"], index=0, key=k("bc", "tone"))

        use_rag = False
        k_chunks = 5
        if RAG_OK:
            with st.expander("Advanced: Knowledge Index (RAG)"):
                use_rag = st.checkbox("Include Knowledge Index", value=False, key=k("bc", "use_rag"))
                k_chunks = st.slider("RAG chunks (top-k)", 1, 7, 5, key=k("bc", "k"))

        if st.button("Generate Business Case Draft", use_container_width=True, key=k("bc", "gen")):
            # keep user on this tab after rerun
            st.session_state["force_tab"] = "Business Case Draft"

            # Build context from phase artifacts
            ctx_parts = [ctx_text]

            # Optionally add RAG chunks
            if RAG_OK and use_rag:
                try:
                    chunks = retrieve(f"Business case inputs for {pid} {target_gate}", k=k_chunks)
                except Exception:
                    chunks = []
                if chunks:
                    rag_txt = "\n\n".join([f"[{t}]\n{snip}" for t, snip in chunks])
                    ctx_parts.append(rag_txt)

            final_context = "\n\n---\n\n".join([c for c in ctx_parts if c])

            system = (
                "You are DecisionMate AI. Produce an executive-grade business case draft as instructed. "
                f"Creativity level: {st.session_state.get(k('sidebar','temp'), 0.2)}."
            )
            prompt = build_business_case_prompt(
                title=bc_title,
                audience=audience,
                industry=industry,
                gate=target_gate,
                tone=tone,
                context=final_context,
            )
            draft = llm(prompt, system)

            if not draft or "error" in str(draft).lower():
                st.error(draft or "No response.")
            else:
                st.success("Draft generated")
                st.markdown(draft)
                st.session_state[k("bc", "draft_text")] = draft

                _save_artifact_button(
                    "AI_Business_Case_Draft",
                    {"title": bc_title, "text": draft, "audience": audience, "gate": target_gate, "tone": tone},
                    pid,
                    ph,
                    key=k("bc", "save_art"),
                )

                if DOCX_OK:
                    data = _docx_bytes(bc_title or "Business Case", draft)
                    if data:
                        st.download_button(
                            "Download DOCX",
                            data=data,
                            file_name="business_case_draft.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=k("bc", "dl"),
                        )
                    else:
                        st.info("DOCX export not available on this system.")

    # --------------------------
    # 5) Guided Business Case (Q&A)
    # --------------------------
    with tabs[4]:
        st.subheader("Guided Business Case (Q&A)")

        g_industry = st.session_state.get("industry") or "general"
        g_default_gate = st.session_state.get("fel_stage", "FEL1")

        with st.form(k("wiz", "form"), clear_on_submit=False):
            st.markdown("**Step 1 — Basics**")
            g_title = st.text_input("Project Title", value=f"Project {pid}", key=k("wiz", "title"))
            g_audience = st.text_input("Audience", value="Executive Steering Committee", key=k("wiz", "audience"))
            g_gate = st.text_input("Target Gate/Stage", value=g_default_gate, key=k("wiz", "gate"))
            g_tone = st.selectbox("Tone", ["concise, executive", "neutral, formal", "persuasive"], index=0, key=k("wiz", "tone"))

            st.markdown("---\n**Step 2 — Problem & Options**")
            g_problem = st.text_area("Problem Statement", placeholder="What problem are we solving? Why now?", key=k("wiz", "problem"))
            g_options = st.text_area(
                "Options Considered (2–3, with pros/cons)",
                placeholder="- Option A: ... Pros ... Cons ...\n- Option B: ...",
                key=k("wiz", "options")
            )

            st.markdown("---\n**Step 3 — Schedule & Economics**")
            g_milestones = st.text_area(
                "Key Milestones (one per line)",
                placeholder="- FEL1: 2025-10-15\n- FEL2: 2026-01-30\n- FID: 2026-06-15\n- First Oil: 2028-03-01",
                key=k("wiz", "milestones")
            )
            g_capex = st.text_input("CAPEX range (e.g., $120–150M, main drivers)", key=k("wiz", "capex"))
            g_opex = st.text_input("OPEX range (e.g., $8–12M/yr, main drivers)", key=k("wiz", "opex"))
            g_econ = st.text_area(
                "Economic Assumptions (pricing, throughput, discount rate, etc.)",
                placeholder="- Brent: $75/bbl\n- Throughput: 45 kbpd\n- WACC: 10%",
                key=k("wiz", "econ")
            )

            st.markdown("---\n**Step 4 — Risks & KPIs**")
            g_risks = st.text_area(
                "Top Risks (owner & mitigation)",
                placeholder="- Risk: ... Owner: ... Mitigation: ...\n- Risk: ...",
                key=k("wiz", "risks")
            )
            g_kpis = st.text_area(
                "KPIs & Success Criteria (5–7)",
                placeholder="- NPV > $X\n- Schedule slip < Y%\n- TRIR < Z",
                key=k("wiz", "kpis")
            )

            st.markdown("---\n**Step 5 — Optional uploads**")
            g_uploads = st.file_uploader(
                "Attach docs to include as context (TXT, CSV, XLSX, DOCX, PDF)",
                type=["txt", "csv", "xlsx", "xls", "docx", "pdf"],
                accept_multiple_files=True,
                key=k("wiz", "uploads"),
            )

            submitted = st.form_submit_button(
                "Generate Draft from Answers + Artifacts + Uploads",
                use_container_width=True
            )

        if submitted:
            # keep user on this tab after rerun
            st.session_state["force_tab"] = "Guided Business Case (Q&A)"

            context_parts = [ctx_text]

            qa_context = f"""
[Q&A Basics]
Title: {g_title}
Audience: {g_audience}
Industry: {g_industry}
Target Gate: {g_gate}
Tone: {g_tone}

[Problem & Options]
Problem: {g_problem}
Options: {g_options}

[Schedule & Economics]
Milestones:
{g_milestones}

CAPEX: {g_capex}
OPEX: {g_opex}
Economic Assumptions:
{g_econ}

[Risks & KPIs]
Risks:
{g_risks}

KPIs:
{g_kpis}
"""
            context_parts.append(qa_context)

            up_texts = read_uploaded_files(g_uploads)
            if up_texts:
                upload_ctx = "\n\n".join([f"[Upload:{name}]\n{text[:8000]}" for name, text in up_texts])
                context_parts.append(upload_ctx)

            final_context = "\n\n---\n\n".join([c for c in context_parts if c and c.strip()])

            system = "You are DecisionMate AI. Produce an executive-grade business case. Use provided context faithfully."
            prompt = build_business_case_prompt(
                title=g_title,
                audience=g_audience,
                industry=g_industry,
                gate=g_gate,
                tone=g_tone,
                context=final_context,
            )
            draft = llm(prompt, system)

            if not draft or "error" in str(draft).lower():
                st.error(draft or "No response from model.")
            else:
                st.success("Draft generated. You can edit below before saving.")
                edited = st.text_area("Editable Draft", value=draft, height=600, key=k("wiz", "edited"))

                _save_artifact_button(
                    "AI_Business_Case_Draft",
                    {
                        "title": g_title,
                        "text": edited,
                        "audience": g_audience,
                        "gate": g_gate,
                        "inputs": {
                            "problem": g_problem,
                            "options": g_options,
                            "milestones": g_milestones,
                            "capex": g_capex,
                            "opex": g_opex,
                            "economics": g_econ,
                            "risks": g_risks,
                            "kpis": g_kpis,
                        },
                    },
                    pid,
                    ph,
                    key=k("wiz", "save_art"),
                )

                if DOCX_OK:
                    data = _docx_bytes(g_title or "Business Case", edited)
                    if data:
                        st.download_button(
                            "Download DOCX",
                            data=data,
                            file_name="business_case_draft.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=k("wiz", "dl"),
                        )
                    else:
                        st.info("DOCX export not available on this system.")
    with tabs[5]:
        ai_optimizer.render()
    
    with tabs[6]:
        ai_pm_analytics.render()
    with tabs[7]:
        ai_benchmarking.render()