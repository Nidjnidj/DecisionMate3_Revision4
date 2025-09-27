from typing import Dict, List
import streamlit as st

def get_project_ids():
    return [
        st.session_state.get("current_project_id") or "P-DEMO",
    ]

def collect_phase_context(project_id: str, phase_id: str) -> Dict:
    """Collect latest artifacts for the active phase into a compact dict."""
    try:
        from artifact_registry import get_latest, list_required_artifacts
    except Exception:
        get_latest = list_required_artifacts = None

    result = {"project_id": project_id, "phase_id": phase_id, "artifacts": []}
    if not get_latest:
        return result

    # Try both OPS and FEL requirements
    reqs: List[Dict] = []
    try:
        reqs = list_required_artifacts(st.session_state.get("current_phase_id", "FEL1")) or []
    except Exception:
        pass

    # also include a few common types if no registry list
    common = ["Reservoir_Profiles", "Well_Plan", "PFD_Package",
              "Equipment_List", "Schedule_Network", "Cost_Model", "Risk_Register"]
    seen = set()
    for r in reqs:
        seen.add(r.get("type"))
    for c in common:
        if c not in seen:
            reqs.append({"workstream": "—", "type": c})

    for r in reqs:
        t = r.get("type")
        rec = None
        try:
            rec = get_latest(project_id, t, phase_id)
        except Exception:
            pass
        if rec:
            result["artifacts"].append({
                "type": t,
                "status": rec.get("status"),
                "workstream": r.get("workstream"),
                "data": rec.get("data", {}),
            })
    return result

def format_context_text(ctx: Dict) -> str:
    lines = [f"Project: {ctx.get('project_id')}", f"Phase: {ctx.get('phase_id')}"]
    for a in ctx.get("artifacts", []):
        lines.append(f"- {a['type']} [{a.get('status','?')}]: {a.get('data')}")
    return "\n".join(lines[:1500])  # keep prompt compact
# --- Simple upload parsers (TXT/CSV/XLSX/DOCX/PDF) ---

import io
from typing import List, Tuple

def read_uploaded_files(files) -> List[Tuple[str, str]]:
    """
    Returns list of (name, extracted_text).
    Non-fatal: if a type can't be parsed, it returns a short notice.
    """
    out = []
    if not files:
        return out
    for f in files:
        name = getattr(f, "name", "upload")
        try:
            text = _extract_text_from_file(f, name)
        except Exception as e:
            text = f"[Could not parse {name}: {e}]"
        out.append((name, text))
    return out

def _extract_text_from_file(file_obj, name: str) -> str:
    lower = name.lower()
    data = file_obj.read()
    file_obj.seek(0)
    if lower.endswith(".txt"):
        return data.decode(errors="ignore")
    if lower.endswith(".csv"):
        import pandas as pd
        df = pd.read_csv(io.BytesIO(data))
        return df.to_string(index=False)
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        import pandas as pd
        df = pd.read_excel(io.BytesIO(data))
        return df.to_string(index=False)
    if lower.endswith(".docx"):
        try:
            from docx import Document
        except Exception:
            return "[python-docx not installed]"
        doc = Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    if lower.endswith(".pdf"):
        try:
            import pdfplumber
        except Exception:
            return "[pdfplumber not installed]"
        text_chunks = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text_chunks.append(page.extract_text() or "")
        return "\n".join(text_chunks)
    # fallback: try bytes->text
    try:
        return data.decode(errors="ignore")
    except Exception:
        return "[Unsupported file type]"
# --- Simple upload parsers (TXT/CSV/XLSX/DOCX/PDF) ---

import io
from typing import List, Tuple

def read_uploaded_files(files) -> List[Tuple[str, str]]:
    """
    Returns a list of (filename, extracted_text).
    If a file can't be parsed, a short note is returned instead of crashing.
    """
    out = []
    if not files:
        return out
    for f in files:
        name = getattr(f, "name", "upload")
        try:
            text = _extract_text_from_file(f, name)
        except Exception as e:
            text = f"[Could not parse {name}: {e}]"
        out.append((name, text))
    return out

def _extract_text_from_file(file_obj, name: str) -> str:
    lower = name.lower()
    data = file_obj.read()
    file_obj.seek(0)
    if lower.endswith(".txt"):
        return data.decode(errors="ignore")
    if lower.endswith(".csv"):
        import pandas as pd
        df = pd.read_csv(io.BytesIO(data))
        return df.to_string(index=False)
    if lower.endswith(".xlsx") or lower.endswith(".xls"):
        import pandas as pd
        df = pd.read_excel(io.BytesIO(data))
        return df.to_string(index=False)
    if lower.endswith(".docx"):
        try:
            from docx import Document
        except Exception:
            return "[python-docx not installed]"
        doc = Document(io.BytesIO(data))
        return "\n".join([p.text for p in doc.paragraphs])
    if lower.endswith(".pdf"):
        try:
            import pdfplumber
        except Exception:
            return "[pdfplumber not installed]"
        text_chunks = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text_chunks.append(page.extract_text() or "")
        return "\n".join(text_chunks)
    # fallback: try bytes->text
    try:
        return data.decode(errors="ignore")
    except Exception:
        return "[Unsupported file type]"
